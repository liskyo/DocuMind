from fastapi import UploadFile, HTTPException
import io
import pypdf
import docx

async def extract_text_from_file(file: UploadFile) -> str:
    filename = file.filename.lower()
    content = await file.read()
    
    try:
        if filename.endswith('.pdf'):
            return extract_pdf(content)
        elif filename.endswith('.docx'):
            return extract_docx(content)
        else:
            # Assume plain text for everything else (txt, md, py, etc.)
            return content.decode('utf-8')
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process file: {str(e)}")

def extract_pdf(content: bytes) -> str:
    try:
        pdf_reader = pypdf.PdfReader(io.BytesIO(content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text if text.strip() else "No text found in PDF (might be an image scan)."
    except Exception as e:
        raise Exception(f"PDF Error: {str(e)}")

def extract_docx(content: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(content))
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        raise Exception(f"DOCX Error: {str(e)}")

def generate_docx(text: str) -> io.BytesIO:
    doc = docx.Document()
    # Simple reconstruction: split by newlines and add paragraphs
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    
    # Save to memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
